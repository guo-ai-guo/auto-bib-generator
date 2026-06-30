import zipfile

from docx import Document

from app import pipeline
from app.models import Job, MatchStatus
from tests.fixtures import library


def _make_docx(path):
    doc = Document()
    doc.add_paragraph("Cats can learn quickly (Doe, 2018).")
    doc.add_paragraph("This claim is unsupported (Nonexistent, 1999).")
    doc.save(path)


def test_end_to_end(tmp_path):
    in_path = tmp_path / "input.docx"
    out_path = tmp_path / "result.docx"
    _make_docx(in_path)

    job = Job(style="apa")
    job.input_path = str(in_path)
    job.output_path = str(out_path)

    pipeline.run_parse_and_match(job, library())

    statuses = {r.status for r in job.results.values()}
    assert MatchStatus.CONFIDENT in statuses   # Doe 2018
    assert MatchStatus.NONE in statuses        # Nonexistent 1999

    stats = pipeline.generate(job)
    assert stats["entries"] == 1
    assert stats["comments_anchored"] >= 1
    assert out_path.exists()

    # Bibliography heading + Doe entry present.
    doc = Document(str(out_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Bibliography" in text
    assert "Doe" in text

    # A real Word comments part was written.
    with zipfile.ZipFile(out_path) as zf:
        assert "word/comments.xml" in zf.namelist()


def test_regenerate_is_idempotent(tmp_path):
    in_path = tmp_path / "input.docx"
    out_path = tmp_path / "result.docx"
    _make_docx(in_path)

    job = Job(style="apa")
    job.input_path = str(in_path)
    job.output_path = str(out_path)
    pipeline.run_parse_and_match(job, library())

    pipeline.generate(job)
    # Re-run using the previous output as input -> should not stack headings.
    job.input_path = str(out_path)
    pipeline.generate(job)

    doc = Document(str(out_path))
    headings = [p.text for p in doc.paragraphs
                if p.style.name.startswith("Heading") and p.text.strip() == "Bibliography"]
    assert len(headings) == 1
