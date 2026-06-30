"""Produce the output .docx: append the bibliography and flag problem citations.

Comments use python-docx 1.2's native `Document.add_comment(runs, ...)`, which
anchors a real Word comment to a run range.

  - Body citations  -> comment anchored to the run(s) covering the citation.
  - Note citations  -> python-docx can't reach inside footnotes/endnotes, so we
    anchor the comment to the note's *reference mark* in the body (the
    superscript number), which sits right where the reader is looking. If the
    reference mark can't be located, the issue is collected into a trailing
    "Citation issues" list so nothing is silently lost.

Idempotent: if a "Bibliography" heading we previously added is present, it (and
everything after it) is removed before regenerating, so re-runs don't stack.
"""
from __future__ import annotations

from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.text.run import Run

from .models import Location

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENT_AUTHOR = "AutoBib"
COMMENT_INITIALS = "AB"
BIB_HEADING = "Bibliography"
_GENERATED_MARK = "Bibliography"  # heading text we recognise on re-run


@dataclass
class Flag:
    location: Location
    container_id: int
    start: int
    end: int
    message: str


def _runs_for_span(paragraph, start: int, end: int) -> list[Run]:
    runs, pos = [], 0
    for run in paragraph.runs:
        rlen = len(run.text)
        if rlen and pos < end and pos + rlen > start:
            runs.append(run)
        pos += rlen
    return runs


def _note_ref_runs(doc) -> dict[tuple[str, int], Run]:
    """Map (location, note_id) -> the body Run carrying that reference mark."""
    out: dict[tuple[str, int], Run] = {}
    for para in doc.paragraphs:
        for run in para.runs:
            el = run._element
            for ref in el.iter(f"{{{W}}}footnoteReference"):
                nid = ref.get(f"{{{W}}}id")
                if nid is not None:
                    out.setdefault((Location.FOOTNOTE.value, int(nid)), run)
            for ref in el.iter(f"{{{W}}}endnoteReference"):
                nid = ref.get(f"{{{W}}}id")
                if nid is not None:
                    out.setdefault((Location.ENDNOTE.value, int(nid)), run)
    return out


def _strip_previous_bibliography(doc) -> None:
    body = doc.element.body
    paras = doc.paragraphs
    cut_from = None
    for i, p in enumerate(paras):
        if p.style.name.startswith("Heading") and p.text.strip() == _GENERATED_MARK:
            cut_from = i
            break
    if cut_from is None:
        return
    for p in paras[cut_from:]:
        p._element.getparent().remove(p._element)


def write_output(input_path: str, output_path: str, *,
                 bib_entries: list[str], flags: list[Flag],
                 engine: str = "citeproc") -> dict:
    doc = Document(input_path)
    _strip_previous_bibliography(doc)

    paras = doc.paragraphs
    ref_runs = _note_ref_runs(doc)
    report: list[str] = []
    anchored = 0

    for flag in flags:
        runs: list[Run] = []
        if flag.location is Location.BODY:
            if 0 <= flag.container_id < len(paras):
                runs = _runs_for_span(paras[flag.container_id], flag.start, flag.end)
                if not runs and paras[flag.container_id].runs:
                    runs = [paras[flag.container_id].runs[0]]
        else:
            run = ref_runs.get((flag.location.value, flag.container_id))
            if run is not None:
                runs = [run]

        if runs:
            try:
                doc.add_comment(runs, text=flag.message,
                                author=COMMENT_AUTHOR, initials=COMMENT_INITIALS)
                anchored += 1
                continue
            except Exception:
                pass
        report.append(flag.message)

    # Bibliography section.
    doc.add_paragraph(BIB_HEADING, style="Heading 1")
    if bib_entries:
        for entry in bib_entries:
            doc.add_paragraph(entry)
    else:
        doc.add_paragraph("(No confidently matched references.)")

    # Trailing report for anything we couldn't anchor as a comment.
    if report:
        doc.add_paragraph("Citation issues", style="Heading 2")
        for msg in report:
            doc.add_paragraph(f"• {msg}")

    doc.save(output_path)
    return {"comments_anchored": anchored, "unanchored": len(report),
            "entries": len(bib_entries), "engine": engine}
