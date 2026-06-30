"""Read the text containers out of a .docx, with stable anchors.

A "container" is a unit of flowing text we scan for citations: a body paragraph,
a footnote, or an endnote. Each carries an id that lets the writer find the same
spot again:
  - BODY:    container_id = index into python-docx's `document.paragraphs`
             (the writer re-opens the same file and indexes identically).
  - FOOTNOTE/ENDNOTE: container_id = the note's numeric w:id.

Footnotes/endnotes are not exposed by python-docx, so we read them straight out
of the docx zip (word/footnotes.xml, word/endnotes.xml) with lxml.
"""
from __future__ import annotations

import zipfile
from dataclasses import dataclass
from typing import Iterator

from docx import Document

from .models import Location

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# Notes Word inserts automatically (separators) use these reserved ids.
_RESERVED_NOTE_IDS = {-1, 0}


@dataclass
class Container:
    location: Location
    container_id: int
    text: str


def _note_text(note_el) -> str:
    """Concatenate all w:t text inside a single footnote/endnote element."""
    parts = []
    for t in note_el.iter(f"{{{W}}}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def _parse_notes(docx_path: str, member: str, location: Location) -> list[Container]:
    from lxml import etree

    containers: list[Container] = []
    with zipfile.ZipFile(docx_path) as zf:
        if member not in zf.namelist():
            return containers
        root = etree.fromstring(zf.read(member))

    tag = "footnote" if location is Location.FOOTNOTE else "endnote"
    for note in root.findall(f"{{{W}}}{tag}"):
        nid = note.get(f"{{{W}}}id")
        if nid is None:
            continue
        nid = int(nid)
        if nid in _RESERVED_NOTE_IDS:
            continue
        text = _note_text(note)
        if text.strip():
            containers.append(Container(location, nid, text))
    return containers


def parse_containers(docx_path: str) -> list[Container]:
    """Return every scannable text container in document order-ish.

    Body paragraphs first (in document order), then footnotes, then endnotes.
    """
    containers: list[Container] = []

    doc = Document(docx_path)
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if text and text.strip():
            containers.append(Container(Location.BODY, i, text))

    containers.extend(_parse_notes(docx_path, "word/footnotes.xml",
                                   Location.FOOTNOTE))
    containers.extend(_parse_notes(docx_path, "word/endnotes.xml",
                                   Location.ENDNOTE))
    return containers


def iter_containers(docx_path: str) -> Iterator[Container]:
    yield from parse_containers(docx_path)
